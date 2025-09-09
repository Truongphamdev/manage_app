from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from django.http import HttpResponse
from ...permissions import IsAdminRole
from ...models import Product
from ...serializers import StockReportSerializer,ReportRevenueSerializer
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, NamedStyle
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
class ReportViewSet(viewsets.ViewSet):
    permission_classes = [IsAdminRole]

    def list(self, request):
        try:
            product_id = request.query_params.get('product_id')
            start_date = request.query_params.get('start_date')
            end_date = request.query_params.get('end_date')
            if not product_id or not start_date or not end_date:
                return Response({"error": "Vui lòng cung cấp product_id, start_date và end_date"}, status=status.HTTP_400_BAD_REQUEST)
            product = Product.objects.filter(ProductID=int(product_id)).first()
            if not product:
                return Response({"error": "Product not found"}, status=status.HTTP_404_NOT_FOUND)
            data = {
                "product_id": int(product_id),
                "start_date": start_date,
                "end_date": end_date
            }
            serializer = StockReportSerializer(data)

            print("Serializer data:", serializer.data)
            return Response(serializer.data)
        except Exception as e:
            print(f"Error in ReportViewSet list: {e}")
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=False, methods=['get'])
    def export_report(self, request):
        product_id = request.query_params.get('product_id')
        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')
        if not product_id or not start_date or not end_date:
            return Response({"error": "Vui lòng cung cấp product_id, start_date và end_date"}, status=status.HTTP_400_BAD_REQUEST)
        product = Product.objects.filter(pk=product_id).first()
        if not product:
            return Response({"error": "Product not found"}, status=404)
        data = {
            "product_id": int(product_id),
            "start_date": start_date,
            "end_date": end_date
        }
        serializer = StockReportSerializer(data)
        # Tạo workbook và sheet
        wb = Workbook()
        ws = wb.active
        ws.title = "Stock Report"

        # ======= Style =======
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill("solid", fgColor="4F81BD")
        center_align = Alignment(horizontal="center", vertical="center")
        right_align = Alignment(horizontal="right", vertical="center")
        thin_border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )

        # ======= Tiêu đề lớn =======
        ws.merge_cells("A1:H1")
        title_cell = ws["A1"]
        title_cell.value = "📊 BÁO CÁO TỒN KHO"
        title_cell.font = Font(size=14, bold=True, color="4F81BD")
        title_cell.alignment = Alignment(horizontal="center", vertical="center")

        # ======= Header =======
        headers = [
            'Mã ID', 'Tên sản phẩm', 'Ngày bắt đầu', 'Ngày kết thúc',
            'Tồn đầu', 'Nhập', 'Xuất', 'Tồn cuối'
        ]
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=2, column=col_num, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center_align
            cell.border = thin_border

        # ======= Data =======
        row_data = [
            serializer.data['product_id'],
            serializer.data['product_name'],
            serializer.data['start_date'],
            serializer.data['end_date'],
            serializer.data['beginning_inventory'],
            serializer.data['imports'],
            serializer.data['exports'],
            serializer.data['ending'],
        ]

        for col_num, value in enumerate(row_data, 1):
            cell = ws.cell(row=3, column=col_num, value=value)
            if isinstance(value, int) or isinstance(value, float):
                cell.alignment = right_align
                # Format số có dấu phân cách
                cell.number_format = '#,##0'
            else:
                cell.alignment = center_align
            cell.border = thin_border

        # ======= Auto fit column width =======
            for col_idx, col in enumerate(ws.columns, 1):
                max_length = 0
                col_letter = get_column_letter(col_idx)
                for cell in col:
                    try:
                        if cell.value:
                            max_length = max(max_length, len(str(cell.value)))
                    except:
                        pass
                ws.column_dimensions[col_letter].width = max_length + 2

        # ======= Response =======
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={'Content-Disposition': f'attachment; filename=stock_report_{product_id}.xlsx'}
        )
        wb.save(response)
        return response

class ReportRevenueViewSet(viewsets.ViewSet):
    permission_classes = [IsAdminRole]

    def list(self, request):
        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')
        if not start_date or not end_date:
            return Response({"error": "Vui lòng cung cấp start_date và end_date"}, status=status.HTTP_400_BAD_REQUEST)
        data = {
            "start_date": start_date,
            "end_date": end_date
        }
        serializer = ReportRevenueSerializer(data)
        return Response(serializer.data)

    @action(detail=False, methods=['get'])
    def export_report(self, request):
        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')
        if not start_date or not end_date:
            return Response({"error": "Vui lòng cung cấp start_date và end_date"}, status=status.HTTP_400_BAD_REQUEST)
        data = {
            "start_date": start_date,
            "end_date": end_date
        }
        serializer = ReportRevenueSerializer(data)
        # Tạo workbook và sheet
        doc = Document()

        # ====== Logo + Tiêu đề ======
        doc.add_paragraph().add_run("CÔNG TY TNHH TrườngPRO").bold = True
        doc.add_paragraph("Địa chỉ: Thành phố Fleiku, Gia Lai").italic = True

        title = doc.add_heading("BÁO CÁO DOANH THU", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Thời gian báo cáo
        time_para = doc.add_paragraph(
            f"Từ ngày {serializer.data['start_date']} đến ngày {serializer.data['end_date']}"
        )
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")  # dòng trống

        # ====== Bảng số liệu ======
        table = doc.add_table(rows=5, cols=2)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Header
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Chỉ tiêu"
        hdr_cells[1].text = "Giá trị (VNĐ)"

        # Dữ liệu
        data_rows = [
            ("Tổng doanh thu", serializer.data["total_revenue"]),
            ("Đã thanh toán", serializer.data["total_paid_purchase"]),
            ("Chưa thanh toán", serializer.data["total_unpaid_purchases"]),
            ("Doanh Thu Ròng", serializer.data["total_amount"]),
        ]

        for i, (label, value) in enumerate(data_rows, start=1):
            row_cells = table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = f"{value:,} VNĐ"

            # Style cho số liệu
            run = row_cells[1].paragraphs[0].runs[0]
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 102, 204)  # xanh dương

        # ====== Nhận xét ======
        doc.add_paragraph("\nNHẬN XÉT:", style="Heading 2")
        doc.add_paragraph(
            "Tổng doanh thu trong kỳ đạt mức khả quan. "
            "Tỷ lệ đơn hàng đã thanh toán chiếm phần lớn. "
            "Cần đẩy mạnh thu hồi công nợ và tối ưu hóa quy trình bán hàng."
        )

        # ====== Người lập báo cáo ======
        doc.add_paragraph("\n\nNgười lập báo cáo", style="Normal").alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.add_paragraph("........................................").alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # ====== Xuất file ======
        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": 'attachment; filename="revenue_report.docx"'},
        )
        doc.save(response)
        return response
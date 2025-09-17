from rest_framework import viewsets,status
from rest_framework.response import Response
from ....permissions import IsAdminRole
from ....serializers import InvoicePurchaseSerializer,ReportInvoicePurchaseSerializer,ReportInvoiceOrderSerializer,InvoiceOrderSerializer
from ....models import InvoicePurchase,InvoiceOrder
from django.http import HttpResponse
from docx import Document
from decimal import Decimal
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from rest_framework.decorators import action
class InvoicePurchaseViewSet(viewsets.ViewSet):
    permission_classes = [IsAdminRole]
    @staticmethod
    def format_currency(value):
        try:
            number = Decimal(value)  # ép về Decimal
            return f"{number:,.0f} ₫"  # làm tròn về số nguyên (không in .00)
        except Exception:
            return str(value)
    def list(self,request):
        invoices = InvoicePurchase.objects.all().order_by('-created_at').prefetch_related('purchase__purchasedetail_set__product')
        serializer = InvoicePurchaseSerializer(invoices, many=True)  # many=True cho list
        return Response(serializer.data)
    def retrieve(self, request, pk=None):
        try:
            invoice = InvoicePurchase.objects.get(pk=pk)
            serializer = ReportInvoicePurchaseSerializer(invoice)
            return Response(serializer.data)
        except InvoicePurchase.DoesNotExist:
            return Response({"detail": "Invoice not found."}, status=status.HTTP_404_NOT_FOUND)
    def destroy(self, request, pk=None):
        try:
            invoice = InvoicePurchase.objects.get(pk=pk)
            invoice.delete()
            return Response({"detail": "Invoice deleted successfully."}, status=status.HTTP_204_NO_CONTENT)
        except InvoicePurchase.DoesNotExist:
            return Response({"detail": "Invoice not found."}, status=status.HTTP_404_NOT_FOUND)

    @action(detail=True, methods=['get'])
    def export(self,request, pk=None):

        try:
            invoice = InvoicePurchase.objects.get(pk=pk)
            serializer = ReportInvoicePurchaseSerializer(invoice)
            data = serializer.data
            doc = Document()

            # Header công ty
            doc.add_paragraph().add_run("CÔNG TY TNHH TrườngPRO").bold = True
            doc.add_paragraph("Địa chỉ: Thành phố Pleiku, Gia Lai").italic = True

            # Tiêu đề
            title = doc.add_heading("HÓA ĐƠN MUA HÀNG", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Thông tin cơ bản
            doc.add_paragraph(f"Số hóa đơn: {data['invoice_number']}")
            doc.add_paragraph(f"Ngày mua: {data['purchase_date'].strftime('%d/%m/%Y') if data['purchase_date'] else 'N/A'}")
            doc.add_paragraph(f"Phương thức: {data['method']}")
            doc.add_paragraph(f"Trạng thái: {data['status']}")
            doc.add_paragraph("")

            # Thông tin nhà cung cấp
            doc.add_heading("Thông tin Nhà cung cấp", level=2)
            doc.add_paragraph(f"Tên: {data['supplier_name']}")
            doc.add_paragraph(f"Email: {data['supplier_email']}")
            doc.add_paragraph(f"SĐT: {data['supplier_phone']}")
            doc.add_paragraph(f"Địa chỉ: {data['supplier_address']}")
            doc.add_paragraph(f"Côn-g ty: {data['supplier_company']}")
            doc.add_paragraph("")

            # Bảng sản phẩm
            doc.add_heading("Danh sách sản phẩm", level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"

            # Header bảng
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Tên sản phẩm"
            hdr_cells[1].text = "Số lượng"
            hdr_cells[2].text = "Giá nhập"
            hdr_cells[3].text = "Thành tiền"

            for item in data['items']:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item["product_name"])
                row_cells[1].text = str(item["quantity"])
                row_cells[2].text = self.format_currency(item['price'])
                row_cells[3].text = self.format_currency(item['total'])

            # Tổng tiền
            doc.add_paragraph("")
            total_p = doc.add_paragraph(f"Tổng số tiền: {data['total_amount']:,} ₫")
            total_p.runs[0].bold = True
            total_p.runs[0].font.size = Pt(14)
            total_p.runs[0].font.color.rgb = RGBColor(200, 0, 0)

            # Người lập hóa đơn
            doc.add_paragraph("\n\nNgười lập hóa đơn").alignment = WD_ALIGN_PARAGRAPH.RIGHT
            doc.add_paragraph(".............................").alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Xuất file Word
            response = HttpResponse(
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": 'attachment; filename="invoice_purchase.docx"'},
            )
            doc.save(response)
            return response
        except InvoicePurchase.DoesNotExist:
            return Response({"detail": "Invoice not found."}, status=status.HTTP_404_NOT_FOUND)
        
class InvoiceOrderViewSet(viewsets.ViewSet):
    permission_classes = [IsAdminRole]
    @staticmethod
    def format_number(value):
        try:
            number = Decimal(value)  # ép về Decimal
            return f"{number:,.0f} ₫"  # làm tròn về số nguyên (không in .00)
        except Exception:
            return str(value)
    def list(self,request):
        invoices = InvoiceOrder.objects.all().order_by('-invoice_date').prefetch_related('order__orderdetail_set__product')
        serializer = InvoiceOrderSerializer(invoices, many=True)
        return Response(serializer.data)
    def retrieve(self, request, pk=None):
        try:
            invoice = InvoiceOrder.objects.get(pk=pk)
            serializer = ReportInvoiceOrderSerializer(invoice)
            return Response(serializer.data)
        except InvoicePurchase.DoesNotExist:
            return Response({"detail": "Invoice not found."}, status=status.HTTP_404_NOT_FOUND)
    @action(detail=True, methods=['get'])
    def export(self,request, pk=None):
        try:
            invoice = InvoiceOrder.objects.get(pk=pk)
            serializer = ReportInvoiceOrderSerializer(invoice)
            data = serializer.data
            doc = Document()

            # Header công ty
            doc.add_paragraph().add_run("CÔNG TY TNHH TrườngPRO").bold = True
            doc.add_paragraph("Địa chỉ: Thành phố Pleiku, Gia Lai").italic = True

            # Tiêu đề
            title = doc.add_heading("HÓA ĐƠN BÁN HÀNG", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Thông tin cơ bản
            doc.add_paragraph(f"Số hóa đơn: {data['invoice_number']}")
            doc.add_paragraph(f"Ngày bán: {data['order_date'].strftime('%d/%m/%Y') if data['order_date'] else 'N/A'}")
            doc.add_paragraph(f"Phương thức: {'Tiền mặt' if data['method'] == 'cash' else 'Chuyển khoản'}")
            doc.add_paragraph(f"Trạng thái: {'Đã trả' if data['status'] == 'paid' else 'Chưa trả'}")
            doc.add_paragraph("")

            # Thông tin khách hàng
            doc.add_heading("Thông tin khách hàng", level=2)
            doc.add_paragraph(f"Tên: {data['customer_name']}")
            doc.add_paragraph(f"Email: {data['customer_email']}")
            doc.add_paragraph(f"SĐT: {data['customer_phone']}")
            doc.add_paragraph(f"Địa chỉ: {data['customer_address']}")
            doc.add_paragraph("")

            # Bảng sản phẩm
            doc.add_heading("Danh sách sản phẩm", level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"

            # Header bảng
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Tên sản phẩm"
            hdr_cells[1].text = "Số lượng"
            hdr_cells[2].text = "Giá bán"
            hdr_cells[3].text = "Thành tiền"

            for item in data['items']:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item["product_name"])
                row_cells[1].text = str(item["quantity"])
                row_cells[2].text = self.format_number(item['price'])
                row_cells[3].text = self.format_number(item['total'])

            # Tổng tiền
            doc.add_paragraph("")
            total_p = doc.add_paragraph(f"Tổng số tiền: {data['total_amount']:,} ₫")
            total_p.runs[0].bold = True
            total_p.runs[0].font.size = Pt(14)
            total_p.runs[0].font.color.rgb = RGBColor(200, 0, 0)

            # Người lập hóa đơn
            doc.add_paragraph("\n\nNgười lập hóa đơn").alignment = WD_ALIGN_PARAGRAPH.RIGHT
            doc.add_paragraph(".............................").alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Xuất file Word
            response = HttpResponse(
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": 'attachment; filename="invoice_purchase.docx"'},
            )
            doc.save(response)
            return response
        except InvoicePurchase.DoesNotExist:
            return Response({"detail": "Invoice not found."}, status=status.HTTP_404_NOT_FOUND)
    def destroy(self, request, pk=None):
        invoice = InvoiceOrder.objects.get(pk=pk)
        invoice.delete()